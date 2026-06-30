#!/usr/bin/env python3
"""Convert KUTUMBA DOCX sources to canonical Markdown with structure preservation."""

import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

REPO_ROOT = Path(r"C:\Development\Workspace\DevotionalRepo\kutumba-family-program")
SRC_DIR = REPO_ROOT / "00-source-materials" / "01-current-kutumba-originals"
EXTRACT_DIR = REPO_ROOT / "build-evidence" / "extraction-working"
PARITY_DIR = REPO_ROOT / "build-evidence" / "extraction-reports"

CANONICAL_MAP = {
    "KUTUMBA Master Operating Model.docx": ("00-foundation", "MASTER-OPERATING-MODEL.md"),
    "KUTUMBA_Governance_Charter_and_Policy_Framework_v1.0.docx": (
        "01-governance", "GOVERNANCE-CHARTER-AND-TEMPLE-RELATIONSHIP.md"
    ),
    "KUTUMBA_Three-Year_Curriculum_Architecture_v1.0.docx": (
        "02-curriculum-architecture", "THREE-YEAR-CURRICULUM-ARCHITECTURE.md"
    ),
    "KUTUMBA_First_Six_Month_Detailed_Curriculum.docx": (
        "03-first-six-months", "FIRST-SIX-MONTH-DETAILED-CURRICULUM.md"
    ),
    "KUTUMBA_Children_Youth_Formation_Operating_Model_v1.0.docx": (
        "04-children-youth", "CHILDREN-YOUTH-FORMATION-OPERATING-MODEL.md"
    ),
    "KUTUMBA_Parent_Formation_and_Family_Care_Operating_Model_v1.0.docx": (
        "05-parent-formation", "PARENT-FORMATION-AND-FAMILY-CARE-OPERATING-MODEL.md"
    ),
    "KUTUMBA_Prasada_Hospitality_Weekly_Operations_Manual.docx": (
        "06-prasadam-operations", "PRASADA-HOSPITALITY-WEEKLY-OPERATIONS-MANUAL.md"
    ),
    "KUTUMBA_Chat_7_Kirtana_Worship_Prayers_Bhakti_Laboratories_Operating_Manual.docx": (
        "07-kirtana-worship-bhakti-labs", "KIRTANA-WORSHIP-PRAYERS-BHAKTI-LABORATORIES-MANUAL.md"
    ),
    "KUTUMBA_Chat_8_Festivals_Yatras_Outdoor_Calendar_Working_Framework.docx": (
        "08-festivals-yatras-calendar", "FESTIVALS-YATRAS-OUTDOOR-CALENDAR-FRAMEWORK.md"
    ),
}


def load_manifest() -> dict:
    yaml_path = REPO_ROOT / "00-source-materials" / "SOURCE-MANIFEST.yaml"
    if yaml_path.exists():
        import yaml
        with open(yaml_path, encoding="utf-8") as f:
            return yaml.safe_load(f)
    with open(REPO_ROOT / "00-source-materials" / "SOURCE-MANIFEST.json", encoding="utf-8") as f:
        return json.load(f)


def heading_level(style_name: str) -> int | None:
    if not style_name:
        return None
    m = re.match(r"Heading\s*(\d+)", style_name, re.I)
    if m:
        return int(m.group(1))
    if style_name.lower() in ("title",):
        return 1
    return None


def para_text(p: Paragraph) -> str:
    parts = []
    for run in p.runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            parts.append(f"***{t}***")
        elif run.bold:
            parts.append(f"**{t}**")
        elif run.italic:
            parts.append(f"*{t}*")
        else:
            parts.append(t)
    return "".join(parts).strip()


def table_to_md(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [para_text(cell.paragraphs[0]) if cell.paragraphs else "" for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    # dedupe merged cells per row
    cleaned = []
    for r in rows:
        seen = []
        for c in r:
            if not seen or c != seen[-1]:
                seen.append(c)
        cleaned.append(seen)
    header = cleaned[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for r in cleaned[1:]:
        while len(r) < len(header):
            r.append("")
        lines.append("| " + " | ".join(r[: len(header)]) + " |")
    return "\n".join(lines)


def iter_block_items(parent):
    from docx.document import Document as DocType
    parent_elm = parent.element.body if isinstance(parent, DocType) else parent._element
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def docx_to_markdown(path: Path) -> tuple[str, dict]:
    doc = Document(path)
    lines = []
    stats = {"paragraphs": 0, "headings": 0, "tables": 0, "chars": 0}
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = para_text(block)
            if not text:
                continue
            stats["paragraphs"] += 1
            lvl = heading_level(block.style.name if block.style else "")
            if lvl:
                stats["headings"] += 1
                lines.append("#" * min(lvl, 6) + " " + text)
            else:
                style = (block.style.name or "").lower()
                if "list bullet" in style:
                    lines.append(f"- {text}")
                elif "list number" in style:
                    lines.append(f"1. {text}")
                else:
                    lines.append(text)
            lines.append("")
        elif isinstance(block, Table):
            md = table_to_md(block)
            if md:
                stats["tables"] += 1
                lines.append(md)
                lines.append("")
    content = "\n".join(lines).strip() + "\n"
    stats["chars"] = len(content)
    return content, stats


def find_source_entry(manifest: dict, filename: str) -> dict | None:
    for e in manifest.get("entries", []):
        if e.get("original_filename") == filename:
            return e
    return None


def add_frontmatter(content: str, entry: dict, rel_dest: str) -> str:
    fm = [
        "---",
        f"source_id: {entry.get('source_id', 'UNKNOWN')}",
        f"source_hash: {entry.get('sha256', '')}",
        f"source_file: {entry.get('original_filename', '')}",
        f"canonical_path: {rel_dest}",
        f"canonicalization_date: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}",
        "status: canonical-working-copy",
        "rights_status: kutumba-authored-project-document",
        "---",
        "",
    ]
    return "\n".join(fm) + content


def main():
    manifest = load_manifest()
    EXTRACT_DIR.mkdir(parents=True, exist_ok=True)
    PARITY_DIR.mkdir(parents=True, exist_ok=True)
    results = []

    docx_files = list(SRC_DIR.rglob("*.docx"))
    for docx_path in sorted(docx_files):
        fname = docx_path.name
        mapping = CANONICAL_MAP.get(fname)
        if not mapping:
            folder, out_name = "00-foundation", fname.replace(".docx", ".md")
        else:
            folder, out_name = mapping

        entry = find_source_entry(manifest, fname) or {}
        try:
            content, stats = docx_to_markdown(docx_path)
        except Exception as exc:
            results.append({"file": fname, "status": "FAILED", "error": str(exc)})
            continue

        rel_dest = f"{folder}/{out_name}"
        dest = REPO_ROOT / folder / out_name
        dest.parent.mkdir(parents=True, exist_ok=True)
        full = add_frontmatter(content, entry, rel_dest)
        dest.write_text(full, encoding="utf-8")

        extract_report = PARITY_DIR / f"{entry.get('source_id', fname)}-extraction.md"
        extract_report.write_text(
            f"# Extraction Report: {fname}\n\n"
            f"- Source ID: {entry.get('source_id', 'N/A')}\n"
            f"- Paragraphs: {stats['paragraphs']}\n"
            f"- Headings: {stats['headings']}\n"
            f"- Tables: {stats['tables']}\n"
            f"- Characters: {stats['chars']}\n"
            f"- Canonical: `{rel_dest}`\n"
            f"- Status: extracted\n"
            f"- Notes: Automated extraction via python-docx; verify tables and formatting manually.\n",
            encoding="utf-8",
        )
        results.append({
            "file": fname,
            "source_id": entry.get("source_id"),
            "canonical": rel_dest,
            "stats": stats,
            "status": "OK",
        })

    parity_path = REPO_ROOT / "build-evidence" / "DOCUMENT-PARITY-REPORT.md"
    with open(parity_path, "w", encoding="utf-8") as f:
        f.write("# Document Parity Report\n\n")
        f.write(f"Generated: {datetime.now(timezone.utc).isoformat()}\n\n")
        f.write("## Summary\n\n")
        ok = sum(1 for r in results if r["status"] == "OK")
        f.write(f"- DOCX files processed: {len(results)}\n")
        f.write(f"- Successful extractions: {ok}\n")
        f.write(f"- Failed: {len(results) - ok}\n\n")
        f.write("## Per-document results\n\n")
        f.write("| Source | Source ID | Canonical | Paragraphs | Headings | Tables | Status |\n")
        f.write("|---|---|---|---|---|---|---|\n")
        for r in results:
            st = r.get("stats", {})
            f.write(
                f"| {r['file']} | {r.get('source_id', '')} | {r.get('canonical', '')} "
                f"| {st.get('paragraphs', '')} | {st.get('headings', '')} "
                f"| {st.get('tables', '')} | {r['status']} |\n"
            )
        f.write("\n## Parity notes\n\n")
        f.write("- Original DOCX files remain authoritative visual baselines in `00-source-materials/01-current-kutumba-originals/`.\n")
        f.write("- Markdown canonical copies are working sources for search, diff, and Git review.\n")
        f.write("- Automated DOCX regeneration is marked as a later controlled capability.\n")
        f.write("- PDF companions (where present) should be cross-checked for any rendering differences.\n")

    print(json.dumps({"processed": len(results), "ok": ok}))
    return 0


if __name__ == "__main__":
    sys.exit(main())
